import os
import json
import re
import logging
import argparse
from docx import Document
from docx.shared import Pt, RGBColor, Inches

# Colors for the DOCX file:
ARTICLE_COLOR = RGBColor(0, 0, 255)    # Blue for article sentences
QUESTION_COLOR = RGBColor(255, 0, 0)     # Red for multi‑choice questions

def replace_curly_quotes(s):
    """Replace curly quotes ‘ and ’ with straight apostrophe '."""
    return s.replace("‘", "'").replace("’", "'")

def process_lesson_json(json_filepath):
    """Load a lesson JSON file and update each multi‑choice question by replacing the blank with the correct answer text."""
    with open(json_filepath, "r", encoding="utf-8") as f:
        lesson = json.load(f)
    
    multi_choice = lesson.get("multi-choice", [])
    choices_answer = lesson.get("choices-answer", [])
    
    # Check that multi-choice list length is exactly 12.
    if len(multi_choice) != 12:
        raise ValueError(f"File {json_filepath}: Expected 12 multi-choice questions but found {len(multi_choice)}.")
    
    # For each question (n from 1 to 12), replace "______" in its question text.
    for i, q in enumerate(multi_choice):
        try:
            correct_letter = choices_answer[i].lower()
        except IndexError:
            raise ValueError(f"File {json_filepath}: Missing correct answer for question {i+1}.")
        correct_text = q.get(correct_letter, "")
        # Replace curly quotes in the correct answer text
        correct_text = replace_curly_quotes(correct_text)
        # Replace the blank marker "______" with the correct answer text.
        updated_question = replace_curly_quotes(q.get("question", "")).replace("______", correct_text)
        q["question"] = updated_question
    lesson["multi-choice"] = multi_choice
    return lesson

def create_lesson_doc(lesson, output_doc_path):
    """
    Create a DOCX file for a lesson.
    The DOCX will contain article sentences (blue) first, followed by multi‑choice questions (red),
    without including the question numbers in the output.
    """
    doc = Document()
    
    # Add article sentences (in blue)
    article_list = lesson.get("article", [])
    for sentence in article_list:
        p = doc.add_paragraph()
        run = p.add_run(replace_curly_quotes(sentence))
        run.font.size = Pt(12)
        run.font.color.rgb = ARTICLE_COLOR

    doc.add_paragraph()  # Blank paragraph between sections
    
    # Add multi-choice questions (in red) without question numbers.
    multi_choice = lesson.get("multi-choice", [])
    for q in multi_choice:
        p = doc.add_paragraph()
        # Use only the question text (without the number)
        q_text = replace_curly_quotes(q.get("question", ""))
        run = p.add_run(q_text)
        run.font.size = Pt(12)
        run.font.color.rgb = QUESTION_COLOR

    doc.save(output_doc_path)

def main():
    parser = argparse.ArgumentParser(
        description="Post-process lesson JSON files: replace blanks in multi-choice questions with correct answer text, then create DOCX files with article sentences (blue) and multi-choice questions (red) without question numbers."
    )
    parser.add_argument("json_dir", help="Directory containing lesson JSON files (lesson_1.json, lesson_2.json, ..., lesson_60.json)")
    parser.add_argument("output_dir", help="Output directory for DOCX files (e.g., Lesson1.docx, Lesson2.docx, etc.)")
    args = parser.parse_args()

    if not os.path.exists(args.output_dir):
        os.makedirs(args.output_dir)

    # Process JSON files for lessons 1 to 60.
    for i in range(1, 61):
        json_filename = os.path.join(args.json_dir, f"lesson_{i}.json")
        if not os.path.exists(json_filename):
            logging.error(f"File not found: {json_filename}")
            continue
        try:
            lesson = process_lesson_json(json_filename)
        except Exception as e:
            logging.error(f"Error processing {json_filename}: {e}")
            continue
        
        # Create a DOCX file for this lesson.
        output_doc = os.path.join(args.output_dir, f"Lesson{i}_post.docx")
        create_lesson_doc(lesson, output_doc)
        logging.info(f"Created DOCX for Lesson {i}: {output_doc}")

if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO)
    main()
