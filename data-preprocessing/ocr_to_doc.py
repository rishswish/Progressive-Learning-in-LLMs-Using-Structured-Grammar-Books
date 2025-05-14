import os
import pytesseract
from PIL import Image
import logging
import argparse
from docx import Document
from docx.shared import Pt, Inches

def ocr_image(image_path, lang="eng+chi_sim"):
    """
    Perform OCR on an image using Tesseract via pytesseract.
    
    :param image_path: Path to the image file.
    :param lang: Languages to use (default: English and Simplified Chinese).
    :return: Recognized text as a string.
    """
    try:
        image = Image.open(image_path)
        text = pytesseract.image_to_string(image, lang=lang)
        return text
    except Exception as e:
        logging.error(f"Error processing {image_path}: {e}")
        return None

def process_text(text):
    """
    Remove completely blank lines from the OCR text while preserving intentional line breaks.
    
    :param text: The raw recognized text.
    :return: Cleaned text with blank lines removed.
    """
    # Split text into lines and filter out lines that are completely empty
    lines = text.splitlines()
    non_blank_lines = [line for line in lines if line.strip() != ""]
    return "\n".join(non_blank_lines)

def create_doc_with_text_pages(input_dir, output_file, lang="eng+chi_sim"):
    """
    Processes images from input_dir, performs OCR on each, cleans the text by removing blank lines
    (while preserving line breaks), and compiles all recognized text into a single Word document.
    Each image's text is placed on its own page with narrow margins and font size 8.
    """
    document = Document()

    # Adjust page margins to narrow (0.5" on each side)
    section = document.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    # Set base font style and size.
    base_font_size = 8
    style = document.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(base_font_size)

    # Supported image file extensions.
    supported_formats = ('.png', '.jpg', '.jpeg', '.tiff', '.bmp')
    image_files = sorted([
        os.path.join(input_dir, f)
        for f in os.listdir(input_dir)
        if f.lower().endswith(supported_formats)
    ])

    total_pages = len(image_files)
    for page_num, image_path in enumerate(image_files, start=1):
        logging.info(f"Processing page {page_num}/{total_pages}: {image_path}")

        # Add a heading with the page number.
        document.add_heading(f"Page {page_num}", level=2)

        # Perform OCR on the image.
        raw_text = ocr_image(image_path, lang=lang)
        if raw_text is None:
            raw_text = "OCR failed for this page."
        
        # Clean the recognized text: remove blank lines but preserve line breaks.
        cleaned_text = process_text(raw_text)

        # Add the cleaned text with font size 8.
        para = document.add_paragraph()
        run = para.add_run(cleaned_text)
        run.font.size = Pt(base_font_size)

        # Insert a page break unless it's the last page.
        if page_num < total_pages:
            document.add_page_break()

    # Save the compiled document.
    document.save(output_file)
    logging.info(f"Document saved to {output_file}")

def main():
    parser = argparse.ArgumentParser(
        description="OCR images to a Word document, placing cleaned text (with preserved line breaks) on one page per image."
    )
    parser.add_argument("input_dir", help="Directory containing image files (one per page).")
    parser.add_argument("output_file", help="Output Word document file (e.g., output.docx).")
    parser.add_argument("--lang", default="eng+chi_sim", 
                        help="Tesseract languages to use (default: eng+chi_sim).")
    args = parser.parse_args()

    create_doc_with_text_pages(args.input_dir, args.output_file, lang=args.lang)

if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO)
    main()